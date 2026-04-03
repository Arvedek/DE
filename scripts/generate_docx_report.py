from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "final_report.md"
OUTPUT = ROOT / "docs" / "final_report_paper.docx"


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("`", "")
    text = text.replace("**", "")
    return text.strip()


def set_default_font(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title_page(document: Document, title: str) -> None:
    title_par = document.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_par.add_run(clean_inline(title))
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Paper-Style Final Report")
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)

    spacer = document.add_paragraph()
    spacer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer.add_run("")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Prepared from the implemented project, warehouse outputs, and dashboard results."
    )
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)

    document.add_section(WD_SECTION.NEW_PAGE)


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            text = clean_inline(value)
            cell_par = table.cell(row_idx, col_idx).paragraphs[0]
            run = cell_par.add_run(text)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(10.5)
            if row_idx == 0:
                run.bold = True


def parse_markdown_to_docx(document: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    in_code_block = False
    code_lang = ""
    buffer: list[str] = []
    table_buffer: list[str] = []
    title_done = False

    def flush_paragraph() -> None:
        nonlocal buffer
        if not buffer:
            return
        text = clean_inline(" ".join(part.strip() for part in buffer if part.strip()))
        if text:
            document.add_paragraph(text)
        buffer = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = []
        for idx, line in enumerate(table_buffer):
            if idx == 1 and re.fullmatch(r"\|\s*[-:| ]+\|", line.strip()):
                continue
            parts = [clean_inline(part) for part in line.strip().strip("|").split("|")]
            rows.append(parts)
        if rows:
            add_table(document, rows)
        table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_paragraph()
            flush_table()
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip("`").strip()
                if code_lang == "mermaid":
                    quote = document.add_paragraph()
                    quote.style = "Intense Quote"
                    quote.add_run("Architecture diagram is represented in the repository source and accompanying dashboard materials.")
            else:
                in_code_block = False
                code_lang = ""
            continue

        if in_code_block:
            if code_lang and code_lang != "mermaid":
                para = document.add_paragraph()
                para.style = "Intense Quote"
                para.add_run(clean_inline(line))
            continue

        if line.startswith("|") and line.endswith("|"):
            flush_paragraph()
            table_buffer.append(line)
            continue
        else:
            flush_table()

        if not line.strip():
            flush_paragraph()
            continue

        if line.startswith("# "):
            if not title_done:
                add_title_page(document, line[2:].strip())
                title_done = True
            else:
                flush_paragraph()
                document.add_heading(clean_inline(line[2:].strip()), level=1)
            continue

        if line.startswith("## "):
            flush_paragraph()
            document.add_heading(clean_inline(line[3:].strip()), level=1)
            continue

        if line.startswith("### "):
            flush_paragraph()
            document.add_heading(clean_inline(line[4:].strip()), level=2)
            continue

        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            item = clean_inline(re.sub(r"^\d+\.\s+", "", line))
            document.add_paragraph(item, style="List Number")
            continue

        if line.startswith("- "):
            flush_paragraph()
            document.add_paragraph(clean_inline(line[2:]), style="List Bullet")
            continue

        buffer.append(line)

    flush_paragraph()
    flush_table()


def main() -> None:
    document = Document()
    set_default_font(document)

    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    parse_markdown_to_docx(document, SOURCE.read_text(encoding="utf-8"))
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
